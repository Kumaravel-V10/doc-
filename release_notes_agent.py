import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os

def create_release_notes(excel_path, output_path):
    print(f"Reading Excel file: {excel_path}")
    xls = pd.ExcelFile(excel_path)
    
    # Read the main sheets
    sheets_to_process = ['Drop1-US', 'User_Stories']
    
    document = Document()
    
    # Title
    title = document.add_heading('Release Notes', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Meta data
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    document.add_paragraph(f"Application Release: Drop 1").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    document.add_heading('1. Introduction', level=1)
    document.add_paragraph("This document outlines the features and user stories included in the Drop 1 release of the application. It consolidates user stories and provides detailed acceptance criteria for each item.")
    
    document.add_heading('2. What\'s New (User Stories)', level=1)
    
    for sheet_name in sheets_to_process:
        if sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name)
            
            document.add_heading(f"Source: {sheet_name}", level=2)
            
            # Map column names to handle both sheets
            id_col = 'Story ID' if 'Story ID' in df.columns else None
            title_col = 'Title' if 'Title' in df.columns else 'Summary' if 'Summary' in df.columns else None
            desc_col = 'UserStory' if 'UserStory' in df.columns else 'Description' if 'Description' in df.columns else None
            ac_col = 'Acceptance Criteria' if 'Acceptance Criteria' in df.columns else None
            feat_col = 'Feature' if 'Feature' in df.columns else None
            
            if not id_col:
                continue
                
            # Fill NAs
            df = df.fillna('')
            
            # Group by Feature if possible
            if feat_col:
                grouped = df.groupby(feat_col)
                for feature, group in grouped:
                    if feature:
                        document.add_heading(f"Feature: {feature}", level=3)
                    
                    for index, row in group.iterrows():
                        story_id = row[id_col]
                        if not story_id:
                            continue
                        
                        story_title = row[title_col] if title_col else ''
                        document.add_heading(f"{story_id}: {story_title}", level=4)
                        
                        if desc_col and row[desc_col]:
                            p = document.add_paragraph()
                            p.add_run("Description/User Story: ").bold = True
                            p.add_run(str(row[desc_col]).strip())
                        
                        if ac_col and row[ac_col]:
                            p = document.add_paragraph()
                            p.add_run("Acceptance Criteria:\n").bold = True
                            p.add_run(str(row[ac_col]).strip())
                        
                        document.add_paragraph("_" * 50)
            else:
                for index, row in df.iterrows():
                    story_id = row[id_col]
                    if not story_id:
                        continue
                        
                    story_title = row[title_col] if title_col else ''
                    document.add_heading(f"{story_id}: {story_title}", level=3)
                    
                    if desc_col and row[desc_col]:
                        p = document.add_paragraph()
                        p.add_run("Description/User Story: ").bold = True
                        p.add_run(str(row[desc_col]).strip())
                    
                    if ac_col and row[ac_col]:
                        p = document.add_paragraph()
                        p.add_run("Acceptance Criteria:\n").bold = True
                        p.add_run(str(row[ac_col]).strip())
                        
                    document.add_paragraph("_" * 50)

    document.save(output_path)
    print(f"Successfully created Release Notes at: {output_path}")

if __name__ == "__main__":
    input_file = r'c:\Users\Agentassist\Downloads\doc-\Drop1-Final Consolidated US 5.xlsx'
    output_file = r'c:\Users\Agentassist\Downloads\doc-\Release_Notes_Template.docx'
    create_release_notes(input_file, output_file)
